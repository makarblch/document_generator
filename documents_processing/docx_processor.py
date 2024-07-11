import re
from docx import Document
from config import REGEX_TAG
from tools.utilities import create_input_path, create_output_path, clear_string, make_set


def remove_empty_paragraphs(doc):
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        p_element = paragraph._element
                        p_element.getparent().remove(p_element)


def open_document(template_path: str):
    try:
        doc = Document(create_input_path(template_path))
        return doc
    except Exception as e:
        print(f"Ошибка при загрузке документа: {e}")
        return None


def replace_in_paragraph(paragraph, values):
    # Сохраняем все начальные атрибуты runs
    runs_attributes = [(run.text, run.bold, run.italic, run.underline, run.font.size, run.font.name) for run in
                       paragraph.runs]

    # Сохраняем текст абзаца и заменяем тэги
    full_text = ''.join([run[0] for run in runs_attributes])
    keys = re.findall(REGEX_TAG, full_text)

    for key in keys:
        if key in values:
            if key in ['month'] and values[key].isdigit():
                full_text = full_text.replace(f'{{{{{key}}}}}', f" {values[key]}. ")
            else:
                full_text = full_text.replace(f'{{{{{key}}}}}', f" {values[key]} ")
            full_text = re.sub(rf'\s*\(.*?\)', '', full_text)

    full_text = re.sub(r'\([^()]*?(Ф\.И\.О|Дата|дата)[^()]*?\)', '', full_text)
    full_text = full_text.replace(' , ', ', ').replace(' . ', '. ')
    full_text = clear_string(full_text)

    # Очищаем абзац и заново добавляем runs с обновленным текстом и исходными атрибутами
    paragraph.clear()
    current_position = 0

    for text, bold, italic, underline, size, name in runs_attributes:
        if current_position >= len(full_text):
            break
        run = paragraph.add_run(full_text[current_position:current_position + len(full_text)])
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.size = size
        run.font.name = name
        current_position += len(full_text)


def fill_docx_template(template_path, output_path, values):
    doc = open_document(template_path)
    if not doc:
        return

    # main body
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, values)

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, values)

    # remove_empty_paragraphs(doc)
    out = create_output_path(output_path)
    doc.save(out)


def get_all_tags(template_path):
    doc = open_document(template_path)
    if not doc:
        return []

    tags = set()

    for paragraph in doc.paragraphs:
        tags.update(re.findall(REGEX_TAG, paragraph.text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    tags.update(re.findall(REGEX_TAG, paragraph.text))

    return tags
